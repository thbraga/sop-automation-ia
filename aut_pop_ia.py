# -*- coding: utf-8 -*-
"""
Created on Tue May 13 13:15:02 2025

# Desenvolvido por Thaina Braga – Projeto de Automação com IA (2025)
# -*- coding: utf-8 -*-
# Atualizado em 13/05/2025 - alterações prompt
# Script Unificado Compactado: Etapas 2, 3 e 4a
"""

# === IMPORTS GERAIS ===

from openai import OpenAI
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph as DocxParagraph
import json
import os
import io
import re
import unicodedata
from collections import defaultdict
from docx.shared import Inches

from io import BytesIO
from PIL import Image, ImageOps, UnidentifiedImageError

# === CONFIGURAÇÕES ===

from dotenv import load_dotenv
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
SCOPE = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive", "https://www.googleapis.com/auth/documents"]
credenciais = ServiceAccountCredentials.from_json_keyfile_name(os.getenv("GOOGLE_CREDENTIAL"), SCOPE)
cliente = gspread.authorize(credenciais)
planilha = cliente.open_by_url("https://docs.google.com/spreadsheets/d/19rUa0dvKCFmTaskwouaBKrzB2kwqmrtedoWzHPujdnQ")
aba = planilha.worksheet("Respostas ao formulário 1")
valores = aba.get_all_values()
cabecalhos = valores[0]
colunas = {h: idx + 1 for idx, h in enumerate(cabecalhos)}


# === FUNÇÕES GERAIS ===

def limpar_nome_arquivo(texto):
    return re.sub(r'[\\/*?:"<>|]', "-", texto.strip())

def extrair_file_id(link):
    padrao = r"(?:id=|/d/)([a-zA-Z0-9_-]{25,})"
    resultado = re.search(padrao, link)
    return resultado.group(1) if resultado else None

def baixar_arquivo_drive(file_id, nome_destino):
    credentials = Credentials.from_service_account_file("credentials.json", scopes=SCOPE)
    drive_service = build("drive", "v3", credentials=credentials)
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.FileIO(nome_destino, "wb")
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()

# --- AUXÍLIOS PARA LER PARÁGRAFOS E TABELAS EM ORDEM ---

def iter_block_items(parent):
    """
    Itera parágrafos e tabelas na ordem em que aparecem no documento.
    parent pode ser doc (Document) ou uma célula (_Cell).
    """
    if isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent._element.body

    for child in parent_elm.iterchildren():
        if child.tag.endswith('}p'):
            yield DocxParagraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

def cell_text(cell: _Cell) -> str:
    """Texto da célula (parágrafos unidos por ' / '). Pode customizar como preferir."""
    parts = []
    for p in cell.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    return " / ".join(parts).strip()

def table_to_markdown(tbl: Table) -> str:
    """
    Converte uma tabela do python-docx em Markdown simples.
    Considera a primeira linha como cabeçalho.
    """
    rows = tbl.rows
    if not rows:
        return ""

    # Linhas -> matriz de strings
    matrix = []
    for r in rows:
        matrix.append([cell_text(c) for c in r.cells])

    # Cabeçalho: primeira linha
    header = matrix[0] if matrix else []
    # Em caso de linhas vazias, evita quebra
    header = [h if h else "" for h in header]

    # Separador Markdown
    sep = ["---" for _ in header]

    # Demais linhas
    body = matrix[1:] if len(matrix) > 1 else []

    # Monta markdown
    md = []
    md.append("| " + " | ".join(header) + " |")
    md.append("| " + " | ".join(sep) + " |")
    for row in body:
        md.append("| " + " | ".join((c if c else "") for c in row) + " |")

    return "\n".join(md).strip()

def _coerce_rgb(img: Image.Image) -> Image.Image:
    """Garante modo compatível (RGB), preservando transparência quando necessário."""
    if img.mode in ("RGBA", "LA"):
        # mantém alfa, mas python-docx lida bem com PNG RGBA
        return img
    if img.mode == "P":  # paleta -> RGB
        return img.convert("RGBA") if "transparency" in img.info else img.convert("RGB")
    if img.mode in ("CMYK", "YCbCr"):
        return img.convert("RGB")
    return img

def bytes_to_png_bytes(src_bytes: bytes) -> bytes:
    """
    Tenta abrir qualquer imagem suportada pelo Pillow e converte para PNG.
    - Corrige orientação EXIF.
    - Converte espaço de cor para RGB/RGBA.
    - Garante PNG válido de verdade.
    """
    bio_in = BytesIO(src_bytes)
    try:
        with Image.open(bio_in) as im:
            # Ajuste de orientação EXIF (se houver)
            try:
                im = ImageOps.exif_transpose(im)
            except Exception:
                pass

            # Converte para modo seguro
            im = _coerce_rgb(im)

            # (opcional) limitar tamanho para evitar estouro de memória
            # im.thumbnail((6000, 6000), Image.LANCZOS)

            bio_out = BytesIO()
            im.save(bio_out, format="PNG", optimize=True)
            return bio_out.getvalue()
    except UnidentifiedImageError:
        # Não é um formato suportado pelo Pillow (ex.: EMF/WMF)
        # -> Devolvemos None para quem chamou decidir o que fazer
        return None
    

# === ETAPA 2: Padronizar POPs com GPT-4o ===

def processar_pop(linha_idx, dados_formulario):
    link_arquivo = dados_formulario.get("Arquivo POP", "")
    file_id = extrair_file_id(link_arquivo)
    if not file_id:
        print(f"⚠️ Linha {linha_idx}: Nenhum ID válido no campo 'Arquivo POP'. Pulando...")
        return

    print(f"🔄 Linha {linha_idx}: Baixando arquivo (ID: {file_id})...")
    baixar_arquivo_drive(file_id, "POP_INPUT.docx")

    doc = Document("POP_INPUT.docx")
    output_folder = "imagens_pop"
    os.makedirs(output_folder, exist_ok=True)

    image_index = 1
    tabela_index = 1
    modified_text = []

    def save_picture_from_rid(r_id) -> str | None:
        """Salva a imagem referenciada pelo rId como PNG real e retorna o nome 'IMAGEM_X.png'."""
        nonlocal image_index
        # Acha a part da imagem a partir do relacionamento do documento
        try:
            image_part = doc.part.related_parts[r_id]
        except KeyError:
            return None

        raw = image_part.blob
        content_type = getattr(image_part, "content_type", "")

        # Converte para PNG real (usa sua função bytes_to_png_bytes do patch anterior)
        png_bytes = bytes_to_png_bytes(raw)
        if png_bytes is None:
            # EMF/WMF ou formato não suportado → ignorar (ou converter com pipeline externo)
            print(f"⚠️ Imagem não suportada ({content_type or 'desconhecido'}) ignorada.")
            return None

        image_name = f"IMAGEM_{image_index}.png"
        with open(os.path.join(output_folder, image_name), "wb") as f:
            f.write(png_bytes)
        image_index += 1
        return image_name

    def extract_images_from_paragraph(p: DocxParagraph) -> str:
        """
        Retorna o texto do parágrafo com marcadores [[IMAGEM_X]] inseridos nos pontos
        onde existiam imagens embutidas.
        """
        out_runs = []
        for r in p.runs:
            r_elm = r._element
            # Procura desenhos (w:drawing) com blip (a:blip) -> r:embed = rId
            blips = r_elm.xpath('.//a:blip')
            if blips:
                # Pode haver texto antes/depois no mesmo run; preserva
                txt = r.text or ""
                # limpa para não duplicar; vamos reconstruir
                if txt:
                    out_runs.append(txt)

                for blip in blips:
                    r_id = blip.get(qn('r:embed'))
                    img_name = save_picture_from_rid(r_id)
                    if img_name:
                        out_runs.append(f"[[{img_name}]]")  # marcador inline
                continue

            # Run normal
            if r.text:
                out_runs.append(r.text)
        # Fallback: se não tinha runs (ou nenhum texto), usa p.text
        if not out_runs and p.text and p.text.strip():
            out_runs.append(p.text.strip())
        return "".join(out_runs).strip()

    def cell_text_with_images(cell: _Cell) -> str:
        """Texto da célula preservando ordem e inserindo marcadores de imagem."""
        parts = []
        for bp in cell.paragraphs:
            t = extract_images_from_paragraph(bp)
            if t:
                parts.append(t)
        return " / ".join(parts).strip()

    def table_to_markdown_with_images(tbl: Table) -> str:
        """Tabela → Markdown simples, preservando marcadores nas células."""
        rows = tbl.rows
        if not rows:
            return ""
        matrix = []
        for r in rows:
            matrix.append([cell_text_with_images(c) for c in r.cells])

        header = [h or "" for h in matrix[0]]
        sep = ["---" for _ in header]
        body = matrix[1:] if len(matrix) > 1 else []

        md = []
        md.append("| " + " | ".join(header) + " |")
        md.append("| " + " | ".join(sep) + " |")
        for row in body:
            md.append("| " + " | ".join((c or "") for c in row) + " |")
        return "\n".join(md).strip()

    # Percorre DOC em ordem: parágrafos e tabelas
    for block in iter_block_items(doc):
        if isinstance(block, DocxParagraph):
            text = extract_images_from_paragraph(block)
            if text:
                modified_text.append(text)

        elif isinstance(block, Table):
            md = table_to_markdown_with_images(block)
            if md:
                modified_text.append(f"[TABELA_{tabela_index}]")
                modified_text.append(md)
                tabela_index += 1

    texto_completo = "\n\n".join(modified_text)


    prompt = f"""
Você receberá o conteúdo bruto de um Procedimento Operacional Padrão (POP).

Sua tarefa é:

- NÃO OMITIR nenhuma informação existente (etapas, campos, observações, imagens).
- REESTRUTURAR o conteúdo de maneira formal, técnica e organizada.
- UTILIZAR verbos no infinitivo nas instruções (ex: iniciar, preencher, concluir).
- CONECTAR as ações usando transições claras (ex: "Após concluir...", "Em seguida...", "Retornar para...").
- EXPANDIR e DETALHAR as etapas: para cada ação, explicar o que deve ser feito, como fazer e qual é a finalidade.
- Sempre que identificar campos a serem preenchidos, apresentar cada item da lista de forma estruturada com:
  - "campo": o nome do campo (em negrito no Word)
  - "descricao": explicação do que é, como preencher e sua importância
- Sempre que identificar transações do SAP, apresentar no formato:
  - **Nome da Transação** explicação do que é, como preencher e sua importância.
- Se o procedimento envolver transações SAP, MENCIONAR o código da transação SAP no objetivo.
- AGRUPAR informações gerais que não sejam específicas de uma atividade em uma seção "Observações Gerais".
- CORRIGIR eventuais erros de ortografia, gramática e digitação.
- MELHORAR a fluidez, eliminando repetições desnecessárias e reorganizando frases, mantendo sempre o sentido original.

IMPORTANTE:
- Estruture cada etapa como uma atividade clara, separando por tópicos (Atividade 1, Atividade 2, etc.).
- Caso uma atividade tenha campos específicos, apresente-os como lista estruturada.
- Se existirem imagens, liste o nome das imagens associadas a cada atividade, mantendo a sequência lógica.
- NÃO invente informações que não existam no conteúdo enviado.
- Estamos utilizando o sheets para inputar essas informações então o número de atividades não deve exceder 10; una atividades compatíveis, se necessário.

Formato obrigatório de resposta, apenas em JSON puro:
{{
  "objetivo": "",
  "atividades": [
    {{
      "nome": "",
      "descricao_texto": "",
      "descricao_lista": [
        {{
          "campo": "",
          "descricao": ""
        }}
      ],
      "imagens": ["IMAGEM_1", "IMAGEM_2"]
    }}
  ],
  "observacoes": [""],
  "analise_melhorias": ""
}}

Conteúdo bruto:
{texto_completo}

⚠️ Retorne apenas o JSON, sem explicações adicionais, sem comentários, sem cabeçalhos extras.
"""

    resposta = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Especialista em padronização de POPs."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2
    )

    resposta_json = resposta.choices[0].message.content.strip()
    resposta_json = resposta_json[resposta_json.find('{'):resposta_json.rfind('}')+1]

    try:
        dados_chatgpt = json.loads(resposta_json)
    except json.JSONDecodeError:
        print(f"❌ Linha {linha_idx}: Erro ao interpretar JSON.")
        return
    
    


    atividades = dados_chatgpt.get("atividades", [])
    observacoes = dados_chatgpt.get("observacoes", [])

    dados_para_inserir = {
        "Objetivo": dados_chatgpt.get("objetivo", ""),
        "Análise de Melhorias": dados_chatgpt.get("analise_melhorias", "")
    }

    for i, atividade in enumerate(atividades):
        n = i + 1
        dados_para_inserir[f"Atividade {n}"] = atividade.get("nome", "")
        dados_para_inserir[f"Descrição {n} (texto)"] = atividade.get("descricao_texto", "")

        # Montar a lista formatada com campo em negrito
        lista_formatada = ""
        lista = atividade.get("descricao_lista", [])
        if isinstance(lista, list) and all(isinstance(item, dict) for item in lista):
            lista_formatada = "\n".join(
                [f"**{item['campo']} –** {item['descricao']}" for item in lista]
            )
        elif isinstance(lista, list):
            lista_formatada = "\n".join(lista)

        dados_para_inserir[f"Descrição {n} (lista)"] = lista_formatada
        dados_para_inserir[f"Imagens {n}"] = ", ".join(atividade.get("imagens", []))

    if observacoes:
        dados_para_inserir["Observações"] = "\n".join(observacoes)

    cabecalhos_atuais = aba.row_values(1)
    colunas = {h: idx + 1 for idx, h in enumerate(cabecalhos_atuais)}

    for cab in dados_para_inserir:
        if cab not in colunas:
            aba.update_cell(1, len(colunas) + 1, cab)
            colunas[cab] = len(colunas) + 1

    for cab, valor in dados_para_inserir.items():
        col = colunas[cab]
        aba.update_cell(linha_idx, col, valor)

    if "Status Padronização" not in colunas:
        aba.update_cell(1, len(colunas) + 1, "Status Padronização")
        colunas["Status Padronização"] = len(colunas) + 1

    aba.update_cell(linha_idx, colunas["Status Padronização"], "Não padronizado")
    print(f"✅ Linha {linha_idx} processada com sucesso.")



# === EXECUTAR ETAPA 2 ===

for i, linha in enumerate(valores[1:], start=2):
    dados = dict(zip(cabecalhos, linha))
    if not dados.get("Objetivo", "").strip() and dados.get("Arquivo POP", "").startswith("http"):
        processar_pop(i, dados)

# === ETAPA 3: Montagem do Documento com Modelo ===

# —— Auxiliares (negrito + imagens + dedup) ——
CAMPOS_NEGRITO = [
    "Qtd. Remessa", "Parc.", "Parceiro", "Detalhe Cabeçalho", "Clientes", "Enter", "Gravar"
]

NEGRITO_PATTERN = re.compile(r'(\*\*.*?\*\*|\'[^\']+\')')
IMG_TOKEN_RE = re.compile(r'\[\[(IMAGEM_\d+\.png)\]\]')

def _unique_preserving_order(seq):
    seen = set()
    out = []
    for x in seq:
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out

def _clean_text_dedup(texto: str) -> str:
    """Remove repetições exatas dentro do mesmo parágrafo (frases idênticas)."""
    if not texto:
        return texto
    # normaliza múltiplos espaços
    txt = re.sub(r'\s+', ' ', texto).strip()
    # quebra por sentenças simples (., !, ?)
    partes = re.split(r'(?<=[\.\!\?])\s+', txt) if txt else []
    # se não formou sentenças, dedup por ' / ' ou mantém
    if len(partes) <= 1:
        # tentar dedup por repetição direta de segmento
        tokens = [t for t in re.split(r'(\s+)', txt) if t]
        dedup = []
        last = None
        for t in tokens:
            if t != last:
                dedup.append(t)
            last = t
        return ''.join(dedup)
    partes = [p.strip() for p in partes if p.strip()]
    partes = _unique_preserving_order(partes)
    return ' '.join(partes)

def add_text_with_markup(par: Paragraph, texto: str):
    """Aplica **...**, '...' e palavras-chave em negrito (sem duplicar conteúdo)."""
    if not texto:
        return
    texto = _clean_text_dedup(texto)
    parts = NEGRITO_PATTERN.split(texto)

    def _run(s, bold=False):
        r = par.add_run(s)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        return r

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _run(part[2:-2], True)
        elif part.startswith("'") and part.endswith("'"):
            _run(part[1:-1], True)
        else:
            # quebra preservando espaços
            for tok in re.findall(r'\s+|[^\s]+', part):
                if tok.isspace():
                    _run(tok, False)
                else:
                    _run(tok, any(ch in tok for ch in CAMPOS_NEGRITO))

def criar_paragrafo_apos(par_ref: Paragraph, texto: str, estilo: str) -> Paragraph:
    novo_par = OxmlElement("w:p")
    par_ref._element.addnext(novo_par)
    par = Paragraph(novo_par, par_ref._parent)
    try:
        par.style = estilo
    except KeyError:
        par.style = "Normal"
    if texto:
        add_text_with_markup(par, texto)
    return par

def criar_lista_personalizada(par_ref: Paragraph, texto: str) -> Paragraph:
    novo_par = OxmlElement("w:p")
    par_ref._element.addnext(novo_par)
    par = Paragraph(novo_par, par_ref._parent)
    try:
        par.style = "Modelomarcadores1"
    except KeyError:
        par.style = "Normal"
    if texto:
        add_text_with_markup(par, texto)
    return par

def encontrar_imagem(caminho_diretorio, nome_base):
    for arquivo in os.listdir(caminho_diretorio):
        if arquivo.lower().startswith(nome_base.lower().split('.')[0]):
            return os.path.join(caminho_diretorio, arquivo)
    return None

def inserir_imagem_apos(par_ref: Paragraph, nome_imagem: str) -> Paragraph:
    caminho_diretorio = "imagens_pop"
    caminho_encontrado = encontrar_imagem(caminho_diretorio, nome_imagem)
    if not caminho_encontrado:
        print(f"⚠️ Imagem não encontrada: {nome_imagem}")
        return par_ref
    par_ref = criar_paragrafo_apos(par_ref, "", estilo="Normal")
    novo_par = OxmlElement("w:p")
    par_ref._element.addnext(novo_par)
    par = Paragraph(novo_par, par_ref._parent)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par.add_run().add_picture(caminho_encontrado, width=Cm(16.47), height=Cm(9.26))
    par = criar_paragrafo_apos(par, "", estilo="Normal")
    return par

def escrever_texto_com_imagens(par_ref: Paragraph, texto: str, estilo="Normal") -> Paragraph:
    """Insere texto (com marcações) e imagens nos tokens [[IMAGEM_X.png]]."""
    if not texto:
        return par_ref
    # dedup em nível de parágrafo
    texto = _clean_text_dedup(texto)
    linhas = texto.splitlines() or [""]
    for linha in linhas:
        cursor_par = criar_paragrafo_apos(par_ref, "", estilo=estilo)
        pos = 0
        for m in IMG_TOKEN_RE.finditer(linha):
            antes = linha[pos:m.start()]
            if antes:
                add_text_with_markup(cursor_par, antes)
            img_name = m.group(1)
            cursor_par = inserir_imagem_apos(cursor_par, img_name)
            cursor_par = criar_paragrafo_apos(cursor_par, "", estilo=estilo)
            pos = m.end()
        resto = linha[pos:]
        if resto:
            add_text_with_markup(cursor_par, resto)
        par_ref = cursor_par
    return par_ref

def escrever_item_lista_com_imagens(par_ref: Paragraph, texto: str) -> Paragraph:
    """Item de lista com tokens de imagem e dedup de conteúdo."""
    if not texto:
        return par_ref
    texto = _clean_text_dedup(texto)
    linhas = texto.splitlines() or [""]
    for linha in linhas:
        novo_par = OxmlElement("w:p")
        par_ref._element.addnext(novo_par)
        par = Paragraph(novo_par, par_ref._parent)
        try:
            par.style = "Modelomarcadores1"
        except KeyError:
            par.style = "Normal"
        pos = 0
        cursor_par = par
        for m in IMG_TOKEN_RE.finditer(linha):
            antes = linha[pos:m.start()]
            if antes:
                add_text_with_markup(cursor_par, antes)
            img_name = m.group(1)
            cursor_par = inserir_imagem_apos(cursor_par, img_name)
            cursor_par = criar_paragrafo_apos(cursor_par, "", estilo="Modelomarcadores1")
            pos = m.end()
        resto = linha[pos:]
        if resto:
            add_text_with_markup(cursor_par, resto)
        par_ref = cursor_par
    return par_ref

# —— Montagem do documento (lendo a linha marcada como "Não padronizado") ——
data_atualizada = aba.get_all_values()
cabecalhos = data_atualizada[0]

try:
    idx_status = cabecalhos.index("Status Padronização")
except ValueError:
    raise Exception("⚠️ Coluna 'Status Padronização' não encontrada.")

for i, linha in enumerate(data_atualizada[1:], start=2):
    status = linha[idx_status].strip().lower() if idx_status < len(linha) else ""
    if status != "não padronizado":
        continue

    dados_dict = dict(zip(cabecalhos, linha))

    # — montar atividades vindas da planilha —
    atividades = []
    j = 1
    while f"Atividade {j}" in cabecalhos:
        idx_atividade = cabecalhos.index(f"Atividade {j}")
        if idx_atividade >= len(linha) or not linha[idx_atividade].strip():
            j += 1
            continue
        try:
            nome = (linha[cabecalhos.index(f"Atividade {j}")].strip())
            descricao = (linha[cabecalhos.index(f"Descrição {j} (texto)")].strip())
            lista_str = (linha[cabecalhos.index(f"Descrição {j} (lista)")].strip())
            imagens_str = (linha[cabecalhos.index(f"Imagens {j}")].strip())
        except IndexError:
            break

        # —— dedup lista e imagens ——
        lista = [s for s in (lista_str.split('\n') if lista_str else []) if s.strip()]
        lista = _unique_preserving_order([_clean_text_dedup(s.strip()) for s in lista])

        imagens = [img.strip() for img in (imagens_str.split(',') if imagens_str else []) if img.strip()]
        imagens = _unique_preserving_order(imagens)

        atividades.append({
            "nome": _clean_text_dedup(nome),
            "descricao": _clean_text_dedup(descricao),
            "lista": lista,
            "imagens": imagens
        })
        j += 1

    observacoes = dados_dict.get("Observações", "").strip().split('\n') if "Observações" in dados_dict else []
    observacoes = _unique_preserving_order([_clean_text_dedup(o.strip()) for o in observacoes if o.strip()])

    # — abrir modelo —
    doc = Document("Modelo de Procedimento Metódico.docx")

    ancora = None
    for par in doc.paragraphs:
        if par.text.strip() == "=== ATIVIDADES AQUI ===":
            ancora = par
            break
    if not ancora:
        raise Exception("❌ Marcador '=== ATIVIDADES AQUI ===' não encontrado no modelo Word.")

    # limpa texto do marcador
    if getattr(ancora, "clear", None):
        ancora.clear()
    else:
        for r in list(ancora.runs):
            r.text = ""

    par_ref = ancora

    # — escrever atividades —
    for atividade in atividades:
        par_ref = criar_paragrafo_apos(par_ref, atividade['nome'], estilo="DescAtividade1")
        par_ref = escrever_texto_com_imagens(par_ref, atividade['descricao'], estilo="Desatividade1")
        for item in atividade.get("lista", []):
            par_ref = escrever_item_lista_com_imagens(par_ref, item)
        from os.path import splitext
        
        desc = atividade.get("descricao", "") or ""
        lista_txt = "\n".join(atividade.get("lista", []))
        
        for imagem in atividade.get("imagens", []):
            base, ext = splitext(imagem)
            nome_sem_ext = base if base else imagem.replace(".png", "")
            token_sem = f"[[{nome_sem_ext}]]"
            token_com = f"[[{nome_sem_ext}.png]]"
            if (token_sem not in desc and token_com not in desc and
                token_sem not in lista_txt and token_com not in lista_txt):
                par_ref = inserir_imagem_apos(par_ref, imagem)


    # — observações —
    if observacoes:
        par_ref = criar_paragrafo_apos(par_ref, "Observações Gerais", estilo="DescAtividade1")
        for obs in observacoes:
            par_ref = criar_lista_personalizada(par_ref, obs)

    # — salvar pré-formatado —
    codigo = limpar_nome_arquivo(dados_dict.get("Código", f"Linha{i}"))
    nome_proc = limpar_nome_arquivo(dados_dict.get("Nome do Procedimento", "Procedimento"))
    nome_doc = f"POP - {codigo} - {nome_proc}.docx"
    doc.save(nome_doc)
    print(f"✅ Documento gerado: {nome_doc}")
    aba.update_cell(i, idx_status + 1, "Pré-formatado")

    
    # === ETAPA 4: Substituição de Placeholders e Geração Final ===

def normalizar(texto):
    if not texto:
        return ""
    texto = unicodedata.normalize('NFD', texto)
    texto = texto.encode('ascii', 'ignore').decode('utf-8')
    return texto.lower().strip()

def substituir_em_paragrafos(paragrafos, dados, mapa_placeholders):
    substituidos = []
    for par in paragrafos:
        texto_original = ''.join(run.text for run in par.runs)
        novo_texto = texto_original
        for placeholder, chave_normalizada in mapa_placeholders.items():
            if placeholder in novo_texto and chave_normalizada in dados:
                novo_texto = novo_texto.replace(placeholder, dados[chave_normalizada])
                substituidos.append((placeholder, dados[chave_normalizada]))
        if novo_texto != texto_original and par.runs:
            par.runs[0].text = novo_texto
            for run in par.runs[1:]:
                run.text = ""
    return substituidos

def substituir_em_tabelas(tabelas, dados, mapa_placeholders):
    substituidos = []
    for tabela in tabelas:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituidos += substituir_em_paragrafos(celula.paragraphs, dados, mapa_placeholders)
                substituidos += substituir_em_tabelas(celula.tables, dados, mapa_placeholders)
    return substituidos

def extrair_placeholders_doc(doc):
    encontrados = set()

    def extrair_em_paragrafos(paragrafos):
        for p in paragrafos:
            encontrados.update(re.findall(r"{{{.*?}}}|{{.*?}}", p.text))

    def extrair_em_tabelas(tabelas):
        for tabela in tabelas:
            for linha in tabela.rows:
                for celula in linha.cells:
                    extrair_em_paragrafos(celula.paragraphs)
                    extrair_em_tabelas(celula.tables)

    extrair_em_paragrafos(doc.paragraphs)
    extrair_em_tabelas(doc.tables)
    for section in doc.sections:
        extrair_em_paragrafos(section.header.paragraphs)
        extrair_em_tabelas(section.header.tables)
        extrair_em_paragrafos(section.footer.paragraphs)
        extrair_em_tabelas(section.footer.tables)

    return encontrados

# Recarregar planilha
valores = aba.get_all_values()
cabecalhos = valores[0]
idx_status = cabecalhos.index("Status Padronização")

# Encontrar linha "Pré-formatado"
linha_alvo = None
for i, linha in enumerate(valores[1:], start=2):
    status = linha[idx_status].strip().lower()
    if status == "pré-formatado":
        linha_alvo = linha
        linha_num = i
        break

if not linha_alvo:
    raise Exception("Nenhuma linha com status 'Pré-formatado' encontrada.")

dados_dict_original = dict(zip(cabecalhos, linha_alvo))
dados_dict = {normalizar(k): v for k, v in dados_dict_original.items()}

mapa_placeholders = {f"{{{{{k.strip()}}}}}": k for k in dados_dict_original}
mapa_placeholders.update({f"{{{{{{{k.strip()}}}}}}}": k for k in dados_dict_original})

codigo = dados_dict_original.get("Código") or dados_dict.get("codigo") or f"Linha{linha_num}"
nome_proc = dados_dict_original.get("Nome do Procedimento") or dados_dict.get("nome do procedimento") or "Procedimento"
versao = dados_dict_original.get("Versão") or dados_dict.get("versao") or "v1"

# Limpar para nome de arquivo
def limpar_nome_arquivo(texto):
    return re.sub(r'[\\/*?:"<>|]', "-", texto.strip())

codigo_limpo = limpar_nome_arquivo(codigo)
nome_proc_limpo = limpar_nome_arquivo(nome_proc)
versao_limpa = limpar_nome_arquivo(versao)

nome_doc = f"POP - {codigo_limpo} - {nome_proc_limpo}.docx"

if not os.path.exists(nome_doc):
    raise FileNotFoundError(f"""
Arquivo não encontrado: {nome_doc}
Dica: Verifique se os campos 'Código' e 'Nome do Procedimento' estão corretamente preenchidos
na aba 'Respostas ao formulário 1' da planilha.
""")


# Abrir documento
doc = Document(nome_doc)

# Substituir Placeholders
placeholders_encontrados = extrair_placeholders_doc(doc)
substituidos = substituir_em_paragrafos(doc.paragraphs, dados_dict_original, mapa_placeholders)
substituidos += substituir_em_tabelas(doc.tables, dados_dict_original, mapa_placeholders)

for section in doc.sections:
    substituidos += substituir_em_paragrafos(section.header.paragraphs, dados_dict_original, mapa_placeholders)
    substituidos += substituir_em_tabelas(section.header.tables, dados_dict_original, mapa_placeholders)
    substituidos += substituir_em_paragrafos(section.footer.paragraphs, dados_dict_original, mapa_placeholders)
    substituidos += substituir_em_tabelas(section.footer.tables, dados_dict_original, mapa_placeholders)

# Debug de placeholders
substituidos_set = {s[0] for s in substituidos}
nao_substituidos = placeholders_encontrados - substituidos_set

if nao_substituidos:
    print("\n⚠️ Placeholders NÃO substituídos:")
    for ph in nao_substituidos:
        print(f"   {ph}")
else:
    print("\n✅ Todos os placeholders foram substituídos com sucesso.")

# Remover trecho entre "=== EXCLUIR ==="
def remover_trecho_para_excluir(doc):
    excluir_inicio = None
    excluir_fim = None
    for i, par in enumerate(doc.paragraphs):
        if '=== EXCLUIR ===' in par.text:
            if excluir_inicio is None:
                excluir_inicio = i
            else:
                excluir_fim = i
                break
    if excluir_inicio is not None and excluir_fim is not None:
        for i in range(excluir_fim, excluir_inicio - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
            p._p = p._element = None
        print("🗑️ Trecho entre '=== EXCLUIR ===' removido.")
    else:
        print("⚠️ Marcadores '=== EXCLUIR ===' não encontrados.")

remover_trecho_para_excluir(doc)

# Salvar novo documento final
nome_final = f"{codigo_limpo} - {nome_proc_limpo} - v.{versao_limpa}.docx"
caminho_destino = os.path.join("resultado", nome_final)
doc.save(caminho_destino)

print(f"\n📄 Documento final salvo em: {caminho_destino}")

# Atualizar status para "Padronizado"
aba.update_cell(linha_num, idx_status + 1, "Padronizado")
# === LIMPEZA DAS IMAGENS ===
imagens_dir = "imagens_pop"
if os.path.exists(imagens_dir):
    for arquivo in os.listdir(imagens_dir):
        caminho_arquivo = os.path.join(imagens_dir, arquivo)
        try:
            if os.path.isfile(caminho_arquivo):
                os.remove(caminho_arquivo)
        except Exception as e:
            print(f"⚠️ Erro ao apagar {arquivo}: {e}")
    print(f"🧹 Imagens da pasta '{imagens_dir}' apagadas com sucesso.")

# === LIMPEZA DE ARQUIVOS INTERMEDIÁRIOS ===
arquivos_intermediarios = ["POP_INPUT.docx", nome_doc]
for arquivo in arquivos_intermediarios:
    try:
        if os.path.exists(arquivo):
            os.remove(arquivo)
            print(f"🗑️ Arquivo intermediário removido: {arquivo}")
    except Exception as e:
        print(f"⚠️ Erro ao remover arquivo intermediário {arquivo}: {e}")
else:
    print(f"⚠️ Pasta '{imagens_dir}' não encontrada.")

